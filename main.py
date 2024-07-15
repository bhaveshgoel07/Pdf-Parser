
from PIL import Image
from pytesseract import image_to_string
import io
from docx import Document
import re
import pandas as pd
import cv2
from PyPDF2 import PdfReader
import numpy as np

df = pd.read_excel("anskey.xlsx")
ans_key = df.set_index('Question ID')['Correct Option ID'].to_dict()

reader = PdfReader("test.pdf")
doc = Document()
doc.add_heading('CUET Questions', 0)


count = 1
for page in reader.pages:
    page_text = page.extract_text()
    pattern = r'Question ID : (\d+)'
    match = re.search(pattern, page_text)
    
    question_id = match.group(1)
    correct_option = 0
    
    for i in range(1,5):
        option_pattern = r'Option '+str(i)+' ID : (\d+)'
        option_id = re.search(option_pattern, page_text).group(1)
        
        correct_option_id = ans_key[int(question_id)]
        
        if int(correct_option_id) == int(option_id.strip()):
            correct_option = i
            break
    
    paragraph = doc.add_paragraph(f"Q.{count}) ", style='List Number')
    for image_file_object in page.images:
        # pil_image = Image.frombytes(image_file_object.data)
        # with open(str(count) + image_file_object.name, "wb") as fp:
        #     fp.write(image_file_object.data)
        #     count += 1
        img_pil = Image.open(io.BytesIO(image_file_object.data))
        img = np.array(img_pil)
        img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        kernel = np.ones((1, 1), np.uint8)
        img = cv2.dilate(img, kernel, iterations=1)
        img = cv2.erode(img, kernel, iterations=1)
        cv2.threshold(cv2.GaussianBlur(img, (5, 5), 0), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

        cv2.threshold(cv2.bilateralFilter(img, 5, 75, 75), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

        cv2.threshold(cv2.medianBlur(img, 3), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

        cv2.adaptiveThreshold(cv2.GaussianBlur(img, (5, 5), 0), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

        cv2.adaptiveThreshold(cv2.bilateralFilter(img, 9, 75, 75), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

        cv2.adaptiveThreshold(cv2.medianBlur(img, 3), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

        text = image_to_string((img),config='--oem 3 --psm 6')
        
        processed_text = re.sub(r'(?<![^\n](\d\.|[A-Z]\.)\s)\n(?!(\d\.|[A-Z]\.)\s)', ' ', text)

        paragraph.add_run(processed_text)

        # run = doc.add_paragraph(processed_text.strip(),style='QuestionNumber')
    paragraph.add_run("\nCorrect Option : "+ str(correct_option)+"\n")
    count +=1

doc.save('my_document7.docx')

