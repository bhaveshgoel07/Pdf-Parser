
from PIL import Image
from pytesseract import image_to_string
import io
from docx import Document
import re
import pandas as pd
import cv2
import numpy as np
from bs4 import BeautifulSoup
import requests
import os
from docx.shared import Pt



def img_to_str(img):
    img_pil = Image.open(io.BytesIO(img))
    img = np.array(img_pil)
    img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    kernel = np.ones((1, 1), np.uint8)
    img = cv2.dilate(img, kernel, iterations=1)
    img = cv2.erode(img, kernel, iterations=1)
    cv2.threshold(cv2.GaussianBlur(img, (5, 5), 0), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

    cv2.threshold(cv2.bilateralFilter(img, 5, 75, 75), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

    cv2.threshold(cv2.medianBlur(img, 3), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

    cv2.adaptiveThreshold(cv2.GaussianBlur(img, (5, 5), 0), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

    cv2.adaptiveThreshold(cv2.bilateralFilter(img, 9, 75, 75), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

    cv2.adaptiveThreshold(cv2.medianBlur(img, 3), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 2)

    text = image_to_string((img),config='--oem 3 --psm 6')
    
    # processed_text = re.sub(r'(?<![^\n](\d\.|[A-Z]\.)\s)\n(?!(\d\.|[A-Z]\.)\s)', '', text)
    pattern = r'(?<![\d\.])\n(?![\dA-Z]\.)'

# Use re.sub() to remove all other newlines
    processed_text = re.sub(pattern, ' ', text)
    return processed_text

df = pd.read_excel("anskey.xlsx")
ans_key = df.set_index('Question ID')['Correct Option ID'].to_dict()

df_cuet = pd.read_excel("CUETFinal.xlsx")



doc = Document()




for i in range(len(df_cuet)):
    url = df_cuet['URL'][i]
    if df_cuet['Type'][i] != 'OCR':
        continue
    doc = Document()
    # Set font size and font type for paragraphs
    style = doc.styles['Normal']
    style.font.name = 'Calibri'  # Set the font type
    style.font.size = Pt(14)  # Set the font size

    # Set font size and font type for characters
    style = doc.styles['Default Paragraph Font']
    style.font.name = 'Calibri'  # Set the font type
    style.font.size = Pt(14)
    print(df_cuet['Name'][i])
    doc_name = str(df_cuet['Name'][i].replace(" ", "_"))+'.docx'

    page = requests.get(url)
    page_content = BeautifulSoup(page.content, "lxml")

    sections = page_content.find_all('div',class_="section-cntnr")
    for section in sections:
        section_name = re.search(r'Section\s*:\s*(\w+\s?\w*)', section.text, re.IGNORECASE).group(1)
        
        questions = section.find_all('table',class_="questionPnlTbl")
        if section_name.strip() != df_cuet['Subject'][i].strip():
            continue
        
        for question in questions:

            pattern = r'Question ID :(\d+)'
            match = re.search(pattern, question.text)
            
            question_id = match.group(1)
            correct_option = 0
            
            for i in range(1,5):
                option_pattern = r'Option '+str(i)+' ID :(\d+)'
                option_id = re.search(option_pattern, question.text).group(1)
                
                correct_option_id = ans_key.get((int(question_id)),0)
                
                if int(correct_option_id) == int(option_id.strip()):
                    correct_option = i
                    break


            question_number = re.findall(r'Q.\d+', question.text)[0]
            imgs = question.find_all("img")
            paragraph = doc.add_paragraph(str(question_number)+") ")
            prefix = "https://cdn3.digialm.com/"
            
            for im in imgs:
                
                image = requests.get(prefix+im['src'])
                para = img_to_str(image.content)
                if para.lower().startswith("match list"):
                    paragraph.add_run("\n")
                    paragraph.add_run().add_picture(io.BytesIO(image.content))
                    
                    continue
                paragraph.add_run(para)
            for i in range(1,5):
                if i == correct_option:
                    doc.add_paragraph(f"*[{i}] {i}")
                    continue
                doc.add_paragraph(f"[{i}] {i}")

            
            doc.add_paragraph("[MARKS] 5")
            doc.add_paragraph("[NEGATIVE MARKS] 1")
            doc.add_paragraph("[TAG] ")
            doc.add_paragraph("[QUESTION TYPE] Multi_choice")
            doc.add_paragraph("[SOLUTION] ")
            doc.add_paragraph("")
    print(doc_name)
    folder_path = 'tosend'

    doc_name = os.path.normpath(doc_name.replace("/", "-"))
    file_path = os.path.join(folder_path, doc_name)
    doc.save(file_path)
